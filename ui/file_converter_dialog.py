"""文件格式转换工具"""
import os
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel,
    QPushButton, QFileDialog, QListWidget, QListWidgetItem,
    QComboBox, QProgressBar, QMessageBox, QGroupBox,
    QFrame, QSizePolicy, QWidget
)
from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QFont, QIcon


class ConvertWorker(QThread):
    """文件转换工作线程"""
    progress = Signal(int)
    finished = Signal(bool, str)

    def __init__(self, input_path, output_format, output_dir):
        super().__init__()
        self.input_path = input_path
        self.output_format = output_format
        self.output_dir = output_dir

    def run(self):
        try:
            # 获取输入文件信息
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            input_ext = os.path.splitext(file_name)[1].lower()

            self.progress.emit(20)

            # 根据输入格式和输出格式进行转换
            if input_ext == '.md' and self.output_format == 'pdf':
                self._convert_md_to_pdf()
            elif input_ext == '.md' and self.output_format == 'word':
                self._convert_md_to_word()
            elif input_ext == '.md' and self.output_format == 'excel':
                self._convert_md_to_excel()
            elif input_ext in ['.docx', '.doc'] and self.output_format == 'pdf':
                self._convert_word_to_pdf()
            elif input_ext in ['.docx', '.doc'] and self.output_format == 'md':
                self._convert_word_to_md()
            elif input_ext in ['.xlsx', '.xls'] and self.output_format == 'pdf':
                self._convert_excel_to_pdf()
            elif input_ext in ['.xlsx', '.xls'] and self.output_format == 'md':
                self._convert_excel_to_md()
            elif input_ext == '.pdf' and self.output_format == 'md':
                self._convert_pdf_to_md()
            elif input_ext == '.pdf' and self.output_format == 'word':
                self._convert_pdf_to_word()
            elif input_ext == '.pdf' and self.output_format == 'excel':
                self._convert_pdf_to_excel()
            else:
                self.finished.emit(False, f"不支持的转换格式: {input_ext} → {self.output_format}")
                return

            self.progress.emit(100)
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.{self.output_format}")
            self.finished.emit(True, f"转换成功!\n输出文件: {output_path}")

        except Exception as e:
            self.finished.emit(False, f"转换失败: {str(e)}")

    def _convert_md_to_pdf(self):
        """Markdown 转 PDF"""
        try:
            import markdown
            from weasyprint import HTML

            self.progress.emit(40)
            with open(self.input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: "Microsoft YaHei", sans-serif; padding: 40px; line-height: 1.6; }}
                    h1, h2, h3 {{ color: #333; }}
                    code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 4px; }}
                    pre {{ background: #f4f4f4; padding: 16px; border-radius: 8px; overflow-x: auto; }}
                    table {{ border-collapse: collapse; width: 100%; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background: #f8f9fa; }}
                </style>
            </head>
            <body>{html_content}</body>
            </html>
            """

            self.progress.emit(70)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.pdf")

            HTML(string=full_html).write_pdf(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install markdown weasyprint")

    def _convert_md_to_word(self):
        """Markdown 转 Word"""
        try:
            import markdown
            from docx import Document

            self.progress.emit(40)
            with open(self.input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            doc = Document()
            lines = md_content.split('\n')

            self.progress.emit(60)
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.docx")

            doc.save(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install python-docx")

    def _convert_md_to_excel(self):
        """Markdown 转 Excel"""
        try:
            import openpyxl

            self.progress.emit(40)
            with open(self.input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Markdown Content"

            lines = md_content.split('\n')
            row = 1
            for line in lines:
                if line.strip():
                    ws.cell(row=row, column=1, value=line)
                    row += 1

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.xlsx")

            wb.save(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install openpyxl")

    def _convert_word_to_pdf(self):
        """Word 转 PDF"""
        try:
            from docx import Document

            self.progress.emit(40)
            doc = Document(self.input_path)

            # 简单转换为文本再转 PDF
            text_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

            import markdown
            from weasyprint import HTML

            self.progress.emit(70)
            html_content = markdown.markdown(text_content)
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head><meta charset="utf-8"></head>
            <body style="font-family: 'Microsoft YaHei', sans-serif; padding: 40px;">
            {html_content}
            </body>
            </html>
            """

            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.pdf")

            HTML(string=full_html).write_pdf(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install python-docx markdown weasyprint")

    def _convert_word_to_md(self):
        """Word 转 Markdown"""
        try:
            from docx import Document

            self.progress.emit(40)
            doc = Document(self.input_path)

            md_lines = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    md_lines.append(f"{'#' * level} {para.text}")
                elif para.style.name == 'List Bullet':
                    md_lines.append(f"- {para.text}")
                else:
                    md_lines.append(para.text)

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.md")

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(md_lines))
        except ImportError:
            raise Exception("请先安装依赖: pip install python-docx")

    def _convert_excel_to_pdf(self):
        """Excel 转 PDF"""
        try:
            import openpyxl

            self.progress.emit(40)
            wb = openpyxl.load_workbook(self.input_path)
            ws = wb.active

            # 生成 HTML 表格
            html_table = "<table border='1'>"
            for row in ws.iter_rows(values_only=True):
                html_table += "<tr>"
                for cell in row:
                    html_table += f"<td>{cell if cell else ''}</td>"
                html_table += "</tr>"
            html_table += "</table>"

            from weasyprint import HTML

            self.progress.emit(70)
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head><meta charset="utf-8"></head>
            <body style="font-family: 'Microsoft YaHei', sans-serif; padding: 20px;">
            {html_table}
            </body>
            </html>
            """

            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.pdf")

            HTML(string=full_html).write_pdf(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install openpyxl weasyprint")

    def _convert_excel_to_md(self):
        """Excel 转 Markdown"""
        try:
            import openpyxl

            self.progress.emit(40)
            wb = openpyxl.load_workbook(self.input_path)
            ws = wb.active

            md_lines = []
            headers = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i == 0:
                    headers = [str(cell) if cell else '' for cell in row]
                    md_lines.append("| " + " | ".join(headers) + " |")
                    md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                else:
                    cells = [str(cell) if cell else '' for cell in row]
                    md_lines.append("| " + " | ".join(cells) + " |")

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.md")

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_lines))
        except ImportError:
            raise Exception("请先安装依赖: pip install openpyxl")

    def _convert_pdf_to_md(self):
        """PDF 转 Markdown"""
        try:
            import PyPDF2

            self.progress.emit(40)
            with open(self.input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_content = ""
                for page in reader.pages:
                    text_content += page.extract_text() + "\n\n"

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.md")

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
        except ImportError:
            raise Exception("请先安装依赖: pip install PyPDF2")

    def _convert_pdf_to_word(self):
        """PDF 转 Word"""
        try:
            import PyPDF2
            from docx import Document

            self.progress.emit(40)
            with open(self.input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                doc = Document()

                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.docx")

            doc.save(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install PyPDF2 python-docx")

    def _convert_pdf_to_excel(self):
        """PDF 转 Excel"""
        try:
            import PyPDF2
            import openpyxl

            self.progress.emit(40)
            with open(self.input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                wb = openpyxl.Workbook()
                ws = wb.active

                row = 1
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        for line in lines:
                            if line.strip():
                                ws.cell(row=row, column=1, value=line)
                                row += 1

            self.progress.emit(80)
            file_name = os.path.basename(self.input_path)
            name_without_ext = os.path.splitext(file_name)[0]
            output_path = os.path.join(self.output_dir, f"{name_without_ext}.xlsx")

            wb.save(output_path)
        except ImportError:
            raise Exception("请先安装依赖: pip install PyPDF2 openpyxl")


class FileConverterDialog(QDialog):
    """文件格式转换工具对话框"""

    # 支持的格式
    FORMATS = {
        'pdf': {'name': 'PDF', 'icon': '📄', 'extensions': ['.pdf']},
        'excel': {'name': 'Excel', 'icon': '📊', 'extensions': ['.xlsx', '.xls']},
        'word': {'name': 'Word', 'icon': '📝', 'extensions': ['.docx', '.doc']},
        'md': {'name': 'Markdown', 'icon': '📋', 'extensions': ['.md']},
    }

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("文件格式转换")
        self.setMinimumSize(750, 520)
        self.selected_files = []
        self._init_ui()

    def _init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(20)

        # 标题
        title = QLabel("文件格式转换")
        title.setStyleSheet("""
            font-size: 20px;
            font-weight: bold;
            color: #1e293b;
            padding-bottom: 8px;
        """)
        layout.addWidget(title)

        # 文件选择区域
        file_group = QGroupBox("选择文件")
        file_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                margin-top: 12px;
                padding: 16px;
                background-color: #fafbff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 16px;
                padding: 0 8px;
                color: #475569;
            }
        """)
        file_layout = QVBoxLayout(file_group)

        # 支持的格式提示
        format_hint = QLabel("支持格式: PDF (.pdf) | Excel (.xlsx, .xls) | Word (.docx, .doc) | Markdown (.md)")
        format_hint.setStyleSheet("color: #64748b; font-size: 12px; margin-bottom: 8px;")
        format_hint.setWordWrap(True)
        file_layout.addWidget(format_hint)

        # 文件列表
        self.file_list = QListWidget()
        self.file_list.setStyleSheet("""
            QListWidget {
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                padding: 8px;
                background-color: white;
                min-height: 100px;
            }
            QListWidget::item {
                padding: 8px 12px;
                border-radius: 6px;
                margin: 2px 0;
            }
            QListWidget::item:selected {
                background-color: #e0e7ff;
            }
            QListWidget::item:hover {
                background-color: #f1f5f9;
            }
        """)
        file_layout.addWidget(self.file_list)

        # 文件操作按钮
        file_btn_layout = QHBoxLayout()

        add_file_btn = QPushButton("📁 添加文件")
        add_file_btn.setStyleSheet("""
            QPushButton {
                background-color: #6366f1;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 20px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #4f46e5;
            }
        """)
        add_file_btn.clicked.connect(self._add_files)
        file_btn_layout.addWidget(add_file_btn)

        add_dir_btn = QPushButton("📂 添加目录")
        add_dir_btn.setStyleSheet("""
            QPushButton {
                background-color: #8b5cf6;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 20px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #7c3aed;
            }
        """)
        add_dir_btn.clicked.connect(self._add_directory)
        file_btn_layout.addWidget(add_dir_btn)

        clear_btn = QPushButton("🗑️ 清空")
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #ef4444;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 20px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #dc2626;
            }
        """)
        clear_btn.clicked.connect(self._clear_files)
        file_btn_layout.addWidget(clear_btn)

        file_btn_layout.addStretch()
        file_layout.addLayout(file_btn_layout)

        layout.addWidget(file_group)

        # 转换设置区域
        settings_widget = QWidget()
        settings_layout = QHBoxLayout(settings_widget)
        settings_layout.setContentsMargins(0, 0, 0, 0)

        # 输出格式
        format_group = QGroupBox("输出格式")
        format_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                padding: 16px;
                background-color: #fafbff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 16px;
                padding: 0 8px;
                color: #475569;
            }
        """)
        format_layout = QVBoxLayout(format_group)

        self.format_combo = QComboBox()
        self.format_combo.addItems(["PDF", "Excel", "Word", "Markdown"])
        self.format_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                padding: 10px 16px;
                font-size: 14px;
                min-width: 150px;
                background-color: white;
            }
            QComboBox:focus {
                border-color: #6366f1;
            }
            QComboBox::drop-down {
                border: none;
                padding-right: 12px;
            }
        """)
        format_layout.addWidget(self.format_combo)
        settings_layout.addWidget(format_group)

        # 输出目录
        output_group = QGroupBox("输出目录")
        output_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                padding: 16px;
                background-color: #fafbff;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 16px;
                padding: 0 8px;
                color: #475569;
            }
        """)
        output_layout = QHBoxLayout(output_group)

        self.output_dir_label = QLabel("未选择目录（将输出到源文件目录）")
        self.output_dir_label.setStyleSheet("color: #64748b; font-size: 12px;")
        self.output_dir_label.setWordWrap(True)
        output_layout.addWidget(self.output_dir_label)

        select_dir_btn = QPushButton("选择目录")
        select_dir_btn.setStyleSheet("""
            QPushButton {
                background-color: #0ea5e9;
                color: white;
                border: none;
                border-radius: 8px;
                padding: 10px 16px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #0284c7;
            }
        """)
        select_dir_btn.clicked.connect(self._select_output_dir)
        output_layout.addWidget(select_dir_btn)

        settings_layout.addWidget(output_group)
        layout.addWidget(settings_widget)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 1px solid #e2e8f0;
                border-radius: 6px;
                text-align: center;
                height: 24px;
                background-color: #f1f5f9;
            }
            QProgressBar::chunk {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #6366f1, stop:1 #8b5cf6);
                border-radius: 5px;
            }
        """)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # 转换按钮
        convert_btn = QPushButton("⚡ 开始转换")
        convert_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #6366f1, stop:1 #8b5cf6);
                color: white;
                border: none;
                border-radius: 10px;
                padding: 14px 28px;
                font-size: 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #4f46e5, stop:1 #7c3aed);
            }
            QPushButton:disabled {
                background-color: #cbd5e1;
            }
        """)
        convert_btn.clicked.connect(self._start_convert)
        layout.addWidget(convert_btn)

        self.output_dir = None
        self.worker = None

    def _add_files(self):
        """添加文件"""
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择文件", "",
            "所有支持格式 (*.pdf *.xlsx *.xls *.docx *.doc *.md);;"
            "PDF 文件 (*.pdf);;"
            "Excel 文件 (*.xlsx *.xls);;"
            "Word 文件 (*.docx *.doc);;"
            "Markdown 文件 (*.md)"
        )
        if files:
            for file in files:
                if file not in self.selected_files:
                    self.selected_files.append(file)
                    file_name = os.path.basename(file)
                    ext = os.path.splitext(file_name)[1].lower()
                    icon = "📄" if ext == ".pdf" else "📊" if ext in [".xlsx", ".xls"] else "📝" if ext in [".docx", ".doc"] else "📋"
                    self.file_list.addItem(f"{icon} {file_name}")

    def _add_directory(self):
        """添加目录中的所有支持文件"""
        dir_path = QFileDialog.getExistingDirectory(self, "选择目录")
        if dir_path:
            supported_ext = ['.pdf', '.xlsx', '.xls', '.docx', '.doc', '.md']
            for root, dirs, files in os.walk(dir_path):
                for file in files:
                    if os.path.splitext(file)[1].lower() in supported_ext:
                        full_path = os.path.join(root, file)
                        if full_path not in self.selected_files:
                            self.selected_files.append(full_path)
                            ext = os.path.splitext(file)[1].lower()
                            icon = "📄" if ext == ".pdf" else "📊" if ext in [".xlsx", ".xls"] else "📝" if ext in [".docx", ".doc"] else "📋"
                            self.file_list.addItem(f"{icon} {file}")

    def _clear_files(self):
        """清空文件列表"""
        self.selected_files.clear()
        self.file_list.clear()

    def _select_output_dir(self):
        """选择输出目录"""
        dir_path = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if dir_path:
            self.output_dir = dir_path
            self.output_dir_label.setText(dir_path)

    def _start_convert(self):
        """开始转换"""
        if not self.selected_files:
            QMessageBox.warning(self, "提示", "请先选择要转换的文件")
            return

        format_map = {"PDF": "pdf", "Excel": "excel", "Word": "word", "Markdown": "md"}
        output_format = format_map[self.format_combo.currentText()]

        # 检查输入格式
        for file in self.selected_files:
            ext = os.path.splitext(file)[1].lower()
            if ext[1:] == output_format or (ext in ['.xlsx', '.xls'] and output_format == 'excel') or (ext in ['.docx', '.doc'] and output_format == 'word'):
                QMessageBox.warning(self, "提示", f"文件 {os.path.basename(file)} 已经是 {self.format_combo.currentText()} 格式")
                return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        # 转换第一个文件（简化实现）
        file = self.selected_files[0]
        output_dir = self.output_dir or os.path.dirname(file)

        self.worker = ConvertWorker(file, output_format, output_dir)
        self.worker.progress.connect(self._on_progress)
        self.worker.finished.connect(self._on_finished)
        self.worker.start()

    def _on_progress(self, value):
        """进度更新"""
        self.progress_bar.setValue(value)

    def _on_finished(self, success, message):
        """转换完成"""
        self.progress_bar.setVisible(False)
        if success:
            QMessageBox.information(self, "完成", message)
        else:
            QMessageBox.warning(self, "失败", message)
